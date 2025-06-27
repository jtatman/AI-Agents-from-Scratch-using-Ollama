# utils/export_manager.py

import json
import csv
import os
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import io

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ExportManager:
    """Handle exporting of agent results to various formats"""
    
    def __init__(self, export_dir: str = "exports"):
        self.export_dir = export_dir
        os.makedirs(export_dir, exist_ok=True)

    def export_to_json(self, data: Dict[str, Any], filename: Optional[str] = None) -> str:
        """Export data to JSON format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.json"
        
        filepath = os.path.join(self.export_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        
        return filepath

    def export_to_csv(self, data: Union[Dict[str, Any], List[Dict]], filename: Optional[str] = None) -> str:
        """Export data to CSV format"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.csv"
        
        filepath = os.path.join(self.export_dir, filename)
        
        # Handle different data structures
        if isinstance(data, dict):
            # Convert dict to list of dicts for CSV
            rows = self._dict_to_csv_rows(data)
        else:
            rows = data
        
        if rows:
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                if rows:
                    writer = csv.DictWriter(f, fieldnames=rows[0].keys())
                    writer.writeheader()
                    writer.writerows(rows)
        
        return filepath

    def export_to_pdf(self, data: Dict[str, Any], title: str = "Agent Results", filename: Optional[str] = None) -> str:
        """Export data to PDF format"""
        if not PDF_AVAILABLE:
            raise ImportError("fpdf2 library not available. Install with: pip install fpdf2")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.pdf"
        
        filepath = os.path.join(self.export_dir, filename)
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=16)
        
        # Title
        pdf.cell(0, 10, title, ln=True, align='C')
        pdf.ln(10)
        
        # Add timestamp
        pdf.set_font("Arial", size=10)
        pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(5)
        
        # Content
        pdf.set_font("Arial", size=12)
        self._add_dict_to_pdf(pdf, data)
        
        pdf.output(filepath)
        return filepath

    def export_to_docx(self, data: Dict[str, Any], title: str = "Agent Results", filename: Optional[str] = None) -> str:
        """Export data to Word document format"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Install with: pip install python-docx")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"export_{timestamp}.docx"
        
        filepath = os.path.join(self.export_dir, filename)
        
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        
        # Add timestamp
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add content
        self._add_dict_to_docx(doc, data)
        
        doc.save(filepath)
        return filepath

    def export_clinical_summary(self, clinical_data: Dict[str, Any], format_type: str = "pdf") -> str:
        """Export clinical summary in a formatted layout"""
        if format_type.lower() == "pdf" and PDF_AVAILABLE:
            return self._export_clinical_pdf(clinical_data)
        elif format_type.lower() == "docx" and DOCX_AVAILABLE:
            return self._export_clinical_docx(clinical_data)
        else:
            return self.export_to_json(clinical_data, "clinical_summary.json")

    def get_export_string(self, data: Dict[str, Any], format_type: str) -> str:
        """Get export data as string for download"""
        if format_type.lower() == "json":
            return json.dumps(data, indent=2, ensure_ascii=False)
        elif format_type.lower() == "csv":
            output = io.StringIO()
            rows = self._dict_to_csv_rows(data)
            if rows:
                writer = csv.DictWriter(output, fieldnames=rows[0].keys())
                writer.writeheader()
                writer.writerows(rows)
            return output.getvalue()
        else:
            return json.dumps(data, indent=2, ensure_ascii=False)

    def _dict_to_csv_rows(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Convert nested dictionary to CSV-friendly format"""
        rows = []
        
        def flatten_dict(d: Dict, prefix: str = "") -> Dict[str, Any]:
            """Flatten nested dictionary"""
            items = []
            for k, v in d.items():
                new_key = f"{prefix}_{k}" if prefix else k
                if isinstance(v, dict):
                    items.extend(flatten_dict(v, new_key).items())
                elif isinstance(v, list):
                    items.append((new_key, "; ".join(map(str, v))))
                else:
                    items.append((new_key, str(v)))
            return dict(items)
        
        flattened = flatten_dict(data)
        rows.append(flattened)
        
        return rows

    def _add_dict_to_pdf(self, pdf: FPDF, data: Dict[str, Any], level: int = 0):
        """Recursively add dictionary content to PDF"""
        for key, value in data.items():
            if level == 0:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, str(key).replace('_', ' ').title(), ln=True)
                pdf.set_font("Arial", size=12)
            else:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 8, "  " * level + str(key).replace('_', ' ').title() + ":", ln=True)
                pdf.set_font("Arial", size=12)
            
            if isinstance(value, dict):
                self._add_dict_to_pdf(pdf, value, level + 1)
            elif isinstance(value, list):
                for item in value:
                    pdf.cell(0, 6, "  " * (level + 1) + "• " + str(item), ln=True)
            else:
                pdf.cell(0, 6, "  " * (level + 1) + str(value), ln=True)
            
            pdf.ln(2)

    def _add_dict_to_docx(self, doc: Document, data: Dict[str, Any], level: int = 0):
        """Recursively add dictionary content to Word document"""
        for key, value in data.items():
            if level == 0:
                heading = doc.add_heading(str(key).replace('_', ' ').title(), level=1)
            else:
                heading = doc.add_heading(str(key).replace('_', ' ').title(), level=min(level + 1, 3))
            
            if isinstance(value, dict):
                self._add_dict_to_docx(doc, value, level + 1)
            elif isinstance(value, list):
                for item in value:
                    p = doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(value))

    def _export_clinical_pdf(self, clinical_data: Dict[str, Any]) -> str:
        """Export clinical data with specialized medical formatting"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"clinical_report_{timestamp}.pdf"
        filepath = os.path.join(self.export_dir, filename)
        
        pdf = FPDF()
        pdf.add_page()
        
        # Header
        pdf.set_font("Arial", "B", 18)
        pdf.cell(0, 15, "Clinical Note Analysis Report", ln=True, align='C')
        pdf.ln(5)
        
        # Metadata
        pdf.set_font("Arial", size=10)
        if 'metadata' in clinical_data:
            metadata = clinical_data['metadata']
            pdf.cell(0, 8, f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
            if 'note_length' in metadata:
                pdf.cell(0, 8, f"Original Note Length: {metadata['note_length']} characters", ln=True)
            if 'confidence_score' in metadata:
                pdf.cell(0, 8, f"Extraction Confidence: {metadata['confidence_score']}", ln=True)
        pdf.ln(10)
        
        # Clinical sections
        sections = ['chief_complaint', 'present_illness', 'medical_history', 'medications', 
                   'allergies', 'vital_signs', 'physical_exam', 'assessment', 'plan']
        
        for section in sections:
            if section in clinical_data and clinical_data[section]:
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, section.replace('_', ' ').title(), ln=True)
                pdf.set_font("Arial", size=12)
                
                if isinstance(clinical_data[section], list):
                    for item in clinical_data[section]:
                        if item and item != "Not mentioned":
                            pdf.cell(0, 6, f"• {item}", ln=True)
                else:
                    pdf.cell(0, 6, str(clinical_data[section]), ln=True)
                pdf.ln(5)
        
        pdf.output(filepath)
        return filepath

    def _export_clinical_docx(self, clinical_data: Dict[str, Any]) -> str:
        """Export clinical data to Word document with medical formatting"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"clinical_report_{timestamp}.docx"
        filepath = os.path.join(self.export_dir, filename)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Clinical Note Analysis Report', 0)
        
        # Metadata
        if 'metadata' in clinical_data:
            metadata = clinical_data['metadata']
            doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            if 'note_length' in metadata:
                doc.add_paragraph(f"Original Note Length: {metadata['note_length']} characters")
            if 'confidence_score' in metadata:
                doc.add_paragraph(f"Extraction Confidence: {metadata['confidence_score']}")
        
        # Clinical sections
        sections = ['chief_complaint', 'present_illness', 'medical_history', 'medications', 
                   'allergies', 'vital_signs', 'physical_exam', 'assessment', 'plan']
        
        for section in sections:
            if section in clinical_data and clinical_data[section]:
                doc.add_heading(section.replace('_', ' ').title(), level=1)
                
                if isinstance(clinical_data[section], list):
                    for item in clinical_data[section]:
                        if item and item != "Not mentioned":
                            doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph(str(clinical_data[section]))
        
        doc.save(filepath)
        return filepath

# Global export manager instance
export_manager = ExportManager()